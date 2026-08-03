from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "AulaRed5_arquitectura_funcional_actual.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "526174"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "EAF6EE"
LIGHT_YELLOW = "FFF7DF"
LIGHT_RED = "FCECEC"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_header_row(row):
    """Mark the first row for screen readers and repeated table headers."""
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table, color="D7DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    ppr.append(node)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "AR Bullet" not in [s.name for s in styles]:
        bullet = styles.add_style("AR Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = styles["AR Bullet"]
    bullet.base_style = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(11)
    bullet.font.color.rgb = RGBColor.from_string(INK)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("AulaRed 5  |  Inventario funcional y arquitectura")
    set_run_font(r, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("Documento de trabajo · 1 de agosto de 2026  |  Página ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(footer)


def add_para(doc, text="", *, bold=False, italic=False, color=INK, size=11, align=None, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_rich_para(doc, parts, *, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    for text, kwargs in parts:
        r = p.add_run(text)
        set_run_font(r, **kwargs)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="AR Bullet")
        if isinstance(item, tuple):
            label, detail = item
            r = p.add_run(label)
            set_run_font(r, bold=True)
            r = p.add_run(detail)
            set_run_font(r)
        else:
            r = p.add_run(item)
            set_run_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)


def add_callout(doc, title, body, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=accent, size="8")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r = p2.add_run(body)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, *, header_fill=LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(str(header))
        set_run_font(r, size=font_size, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_status_table(doc):
    add_table(doc, ["Estado", "Significado", "Uso en este documento"], [
        ("Completo", "Existe interfaz y conexión funcional con el backend.", "Se puede usar hoy, sujeto a pruebas de datos y contenido."),
        ("Parcial", "Existe una parte del flujo, pero falta UI, integración o comportamiento completo.", "Debe entrar en la lista priorizada de cierre."),
        ("Pendiente", "Hay una base técnica, idea o endpoint, pero no hay una experiencia utilizable.", "Requiere definición y trabajo de producto."),
        ("No evaluado", "No pertenece al alcance de esta revisión.", "Seguridad queda deliberadamente fuera por indicación del usuario."),
    ], [1700, 3300, 4360])


def build_document():
    doc = Document()
    style_document(doc)

    # First page: memo masthead.
    add_para(doc, "INVENTARIO FUNCIONAL Y ARQUITECTURA", bold=True, color=BLUE, size=10, after=4)
    add_para(doc, "AulaRed 5", bold=True, color=INK, size=26, after=4)
    add_para(doc, "Documento de revisión para completar la aplicación", color=MUTED, size=14, after=18)
    metadata = [
        ("Fecha de revisión", "1 de agosto de 2026"),
        ("Alcance", "Arquitectura, funcionalidades, opciones, roles y navegación"),
        ("Fuera de alcance", "Seguridad, endurecimiento y cumplimiento normativo"),
        ("Base revisada", "Frontend React/Vite, API FastAPI, SQLite, WebSocket y documentación local"),
    ]
    for label, value in metadata:
        add_rich_para(doc, [(f"{label}: ", {"size": 10.5, "color": INK, "bold": True}), (value, {"size": 10.5, "color": INK})], after=2)
    add_callout(doc, "Conclusión ejecutiva", "La plataforma ya tiene un núcleo funcional de aprendizaje: feed, publicaciones, lecturas, actividades, progreso, creación de quizzes interactivos, moderación docente y salas de juego en vivo. La principal brecha no es la ausencia de backend, sino la diferencia entre capacidades técnicas y experiencias visibles: gestión de cursos/usuarios, creación de salas y comentarios todavía no están completamente expuestos en la interfaz.", fill=LIGHT_BLUE, accent=BLUE)
    add_para(doc, "Documento preparado a partir de la estructura y el código actual del proyecto. No representa una propuesta de rediseño ni una auditoría de seguridad.", italic=True, color=MUTED, size=9.5, after=0)

    doc.add_page_break()
    add_heading(doc, "Contenido del documento", 1)
    add_numbered(doc, [
        "Resumen de cobertura actual",
        "Arquitectura técnica y flujo de datos",
        "Roles y permisos funcionales",
        "Mapa completo de navegación",
        "Inventario de módulos y funcionalidades",
        "Motor unificado de quizzes y juego en vivo",
        "Catálogo de API y contratos principales",
        "Modelo de datos y relaciones",
        "Estado de completitud y brechas",
        "Backlog funcional recomendado",
        "Criterios de aceptación para considerar la aplicación completa",
    ])
    add_callout(doc, "Cómo leerlo", "Las secciones describen lo que existe hoy y separan las capacidades implementadas de las que solo están disponibles como API o estructura de datos. Las recomendaciones al final están ordenadas por impacto funcional, sin entrar en seguridad.", fill=LIGHT_YELLOW, accent="A16207")

    add_heading(doc, "1. Resumen de cobertura actual", 1)
    add_para(doc, "AulaRed 5 es una plataforma educativa para quinto grado centrada en reproducción humana, pubertad, autocuidado y aprendizaje social. Combina contenido editorial, publicaciones de estudiantes, actividades evaluables y un motor de quizzes que permite mezclar preguntas con material multimedia.")
    add_table(doc, ["Área", "Situación actual", "Evaluación"], [
        ("Aprendizaje individual", "Feed, lecturas, actividades, ruta, progreso y perfil.", "Completo en el recorrido principal."),
        ("Creación docente", "Constructor de quizzes con seis tipos de bloque, medios y configuración.", "Completo para crear el quiz base; falta gestión más amplia del ciclo de vida."),
        ("Interacción social", "Publicaciones, filtros, imagen de portada y moderación.", "Completo con comentarios parcialmente conectados."),
        ("Docente", "Dashboard, moderación, progreso y acceso al creador.", "Completo como panel inicial; puede crecer en administración de cursos y contenidos."),
        ("Juego en vivo", "Salas, QR, WebSocket, lobby, respuestas y podio.", "Backend y experiencia de sala presentes; falta pantalla clara para crear/listar salas."),
        ("Administración", "Rol administrador y endpoints de usuarios/cursos/categorías.", "Parcial: hay soporte técnico, no un backoffice completo visible."),
    ], [1900, 4400, 3060])
    add_status_table(doc)

    add_heading(doc, "2. Arquitectura técnica y flujo de datos", 1)
    add_heading(doc, "2.1 Capas del sistema", 2)
    add_table(doc, ["Capa", "Tecnología / ubicación", "Responsabilidad"], [
        ("Cliente", "React 18 + React Router 6 + Vite", "Renderiza rutas, layouts, formularios, estados y experiencias de juego."),
        ("Estilos", "Tailwind CSS + frontend/src/styles/index.css", "Tokens visuales, responsive, tarjetas, navegación y estados."),
        ("Estado global", "AuthContext, DataContext, WebSocketContext", "Sesión, datos compartidos, progreso, pendientes y sala en vivo."),
        ("Servicios cliente", "frontend/src/services/api.js", "Cliente fetch con token Bearer y funciones de dominio."),
        ("API", "FastAPI en api/app.py", "Dispatcher compatible, CORS, healthcheck, uploads, WebSocket y respuestas JSON."),
        ("Dominio", "api/routes + api/services", "Rutas HTTP, validación, reglas de negocio y consultas a SQLite."),
        ("Persistencia", "SQLite blog.db + schema.sql", "Usuarios, contenidos, actividades, quizzes, intentos, recursos, auditoría y salas."),
        ("Tiempo real", "WebSocket /ws/room/{room_code}", "Lobby, inicio, respuestas, resultados y podio de salas."),
        ("Operación local", "start-debug.bat + logs/", "Arranque idempotente de frontend/backend y captura de logs de tráfico."),
    ], [1800, 3300, 4260])
    add_heading(doc, "2.2 Flujo de una petición", 2)
    add_numbered(doc, [
        "Una pantalla React usa una función de frontend/src/services/api.js o una petición fetch directa para llamar a /api.",
        "El proxy de Vite envía la petición al backend en el puerto 8000 durante desarrollo.",
        "api/app.py lee JSON, multipart o query params y entrega la petición al registro de rutas.",
        "La ruta aplica autenticación y rol cuando corresponde, y llama a un servicio de dominio.",
        "El servicio consulta o actualiza SQLite, conserva relaciones y devuelve un objeto serializable.",
        "La pantalla actualiza su estado local o el contexto global y muestra carga, éxito, vacío o error.",
    ])
    add_callout(doc, "Decisión arquitectónica importante", "La aplicación conserva un dispatcher propio basado en expresiones regulares para reutilizar handlers existentes, aunque el servidor sea FastAPI. Esto funciona como una capa de compatibilidad, pero aumenta la necesidad de mantener un catálogo de rutas y contratos bien documentado.", fill=LIGHT_YELLOW, accent="A16207")
    add_heading(doc, "2.3 Estructura de frontend", 2)
    add_bullets(doc, [
        ("App.jsx: ", "declara las dos áreas principales: rutas de estudiante bajo / y rutas de docente/administrador bajo /admin."),
        ("components/layout: ", "StudentLayout, TeacherLayout, TopBar, SideMenu, BottomNav, ProtectedRoute y ErrorBoundary."),
        ("pages: ", "pantallas por dominio: feed, exploración, creación, publicaciones, actividades, juegos, perfil, progreso y moderación."),
        ("engines: ", "GameEngine y UnifiedQuiz; el motor unificado concentra la experiencia interactiva."),
        ("data/mockData.js: ", "catálogos visuales y datos de apoyo; el contenido principal se carga desde la API."),
        ("utils: ", "transformaciones de feed, etiquetas de área, formato de fechas y parseo de contenido."),
    ])

    add_heading(doc, "3. Roles y permisos funcionales", 1)
    add_table(doc, ["Rol", "Entrada", "Puede hacer desde la interfaz", "Limitación observada"], [
        ("Estudiante", "/", "Explorar contenidos, leer, jugar actividades, enviar publicaciones, consultar ruta/progreso/perfil y unirse a salas.", "No dispone de administración de cursos ni de creación de quizzes."),
        ("Docente", "/admin", "Ver panel de grupo, crear quizzes, moderar publicaciones, consultar progreso, publicar contenido y operar una sala.", "La creación/listado de salas no está expuesta de forma completa en el flujo principal."),
        ("Administrador", "/admin", "Comparte el área docente y puede gestionar entidades según los endpoints de rol administrador.", "Falta un backoffice visible para usuarios, cursos, categorías y auditoría."),
    ], [1500, 1300, 4100, 2460])
    add_para(doc, "La autorización de navegación se concentra en ProtectedRoute. El backend también aplica require_role en operaciones docentes, administrativas y de moderación. Este documento no evalúa si esas barreras son suficientes desde el punto de vista de seguridad.", italic=True, color=MUTED, size=9.5)

    add_heading(doc, "4. Mapa completo de navegación", 1)
    add_heading(doc, "4.1 Área de estudiante", 2)
    add_table(doc, ["Ruta", "Pantalla", "Función principal", "Estado"], [
        ("/", "FeedPage", "Inicio social con retos rápidos, publicaciones, filtros y creación.", "Completo"),
        ("/explorar", "ExplorePage", "Buscar y filtrar temas, lecturas, actividades y recursos.", "Completo"),
        ("/crear", "CreatePage", "Wizard de publicación en cinco pasos con imagen opcional.", "Completo"),
        ("/mi-ruta", "MyRoutePage", "Ruta de aprendizaje, unidades, pasos y logros.", "Completo con lógica de estados mejorable"),
        ("/perfil", "ProfilePage", "Identidad, avance, logros, temas, publicaciones y cierre de sesión.", "Completo"),
        ("/juegos", "GameLauncherPage", "Catálogo de quizzes y acceso al motor unificado.", "Completo"),
        ("/sala", "GameJoinPage", "Introducir código para unirse a una sala.", "Completo"),
        ("/sala/:code", "GameRoomPage", "Sala en vivo, QR, lobby, juego y podio.", "Completo cuando existe el código"),
        ("/lecturas", "PublicationsPage", "Listado, búsqueda y filtro por área.", "Completo"),
        ("/lecturas/:slug", "PublicationDetailPage", "Lectura, imagen, contenido y comentarios.", "Parcial: comentario simulado en UI"),
        ("/actividades", "ActivitiesPage", "Listado de actividades.", "Completo"),
        ("/actividades/:id", "ActivityDetailPage", "Descripción, intentos, reproducción y envío de respuestas.", "Completo"),
        ("/actividades/:id/resultado", "ActivityResultPage", "Resultado de intento.", "Completo"),
        ("/recursos", "ResourcesPage", "Filtro por área/tipo y apertura de archivo o URL.", "Completo"),
        ("/progreso", "ProgressPage", "Resumen de avance.", "Completo"),
        ("/ruta", "Redirect", "Compatibilidad histórica: redirige a /mi-ruta.", "Completo"),
    ], [1700, 1900, 4000, 1760], font_size=8.8)
    add_heading(doc, "4.2 Área docente y administrador", 2)
    add_table(doc, ["Ruta", "Pantalla", "Función principal", "Estado"], [
        ("/admin", "AdminDashboard", "Panel docente: prioridades, KPIs, rendimiento, estudiantes en riesgo y acciones.", "Completo"),
        ("/admin/crear-juego", "TeacherCreatorPage", "Constructor de quizzes interactivos.", "Completo"),
        ("/admin/juegos", "GameLauncherPage", "Catálogo de quizzes con enlace de creación.", "Completo"),
        ("/admin/moderar", "ModerationPage", "Aprobar, pedir cambios, rechazar o archivar publicaciones.", "Completo"),
        ("/admin/progreso", "ProgressPage", "Consulta de progreso usando vista compartida.", "Completo"),
        ("/admin/crear", "CreatePage", "Crear una publicación desde contexto docente.", "Completo, aunque el wizard conserva lenguaje de publicación"),
        ("/admin/lecturas", "PublicationsPage", "Consultar lecturas.", "Completo"),
        ("/admin/actividades", "ActivitiesPage", "Consultar actividades.", "Completo"),
        ("/admin/recursos", "ResourcesPage", "Consultar recursos.", "Completo"),
        ("/admin/actividades/:id", "ActivityDetailPage", "Abrir una actividad desde el contexto docente.", "Completo"),
        ("/admin/sala/:code", "GameRoomPage", "Operar una sala docente existente.", "Completo con entrada manual"),
        ("/admin/perfil", "ProfilePage", "Perfil reutilizado.", "Completo, pero no es un perfil docente especializado"),
    ], [1900, 1900, 3900, 1660], font_size=8.8)
    add_callout(doc, "Observación de navegación", "La separación por layout está bien definida: el estudiante recibe una barra inferior y una experiencia de feed; el docente recibe sidebar, panel de clase y accesos de creación. El siguiente salto debería ser hacer visible el ciclo de vida de salas y quizzes desde el panel docente, sin depender de códigos o rutas conocidas.", fill=LIGHT_GREEN, accent="15803D")

    add_heading(doc, "5. Inventario de módulos y funcionalidades", 1)
    add_heading(doc, "5.1 Autenticación y sesión", 2)
    add_bullets(doc, [
        ("Login: ", "formulario por código y contraseña; el destino se decide según el rol."),
        ("Persistencia de sesión: ", "token y usuario se guardan en localStorage y se recuperan al recargar."),
        ("Cierre de sesión: ", "limpia el estado local y navega a /login."),
        ("Usuario actual: ", "API disponible en /api/auth/me, aunque el cliente se apoya principalmente en el usuario almacenado."),
    ])
    add_heading(doc, "5.2 Feed y publicaciones", 2)
    add_bullets(doc, [
        ("Feed: ", "combina publicaciones y actividades como tarjetas, permite filtrar por Todo, Lecturas, Actividades y Compañeros."),
        ("Publicación destacada: ", "se muestra fijada cuando el contenido viene marcado como destacada."),
        ("Creación: ", "wizard de cinco pasos: propósito, tema, contenido, checklist y envío."),
        ("Tipos de publicación: ", "aprendizaje, pregunta, solución, observación, dibujo y evidencia."),
        ("Imagen de portada: ", "JPG/PNG/WEBP hasta 5 MB, con vista previa y carga posterior a la creación."),
        ("Moderación: ", "estados enviada, en revisión, requiere cambios, aprobada, rechazada y archivada."),
    ])
    add_heading(doc, "5.3 Lecturas, temas y recursos", 2)
    add_bullets(doc, [
        ("Explorar: ", "pestañas para temas, lecturas, actividades y recursos; búsqueda textual dentro de cada conjunto."),
        ("Lecturas: ", "filtro por área, búsqueda por título/resumen y detalle con portada, autor y contenido HTML."),
        ("Recursos: ", "filtro por área y tipo; soporta PDF, imagen y video mediante archivo o URL."),
        ("Categorías: ", "áreas temáticas de cuerpo humano, pubertad, sistemas reproductores, células, fecundación, embarazo, nacimiento, autocuidado y preguntas frecuentes."),
    ])
    add_heading(doc, "5.4 Ruta, progreso y logros", 2)
    add_bullets(doc, [
        ("Ruta: ", "agrupa lecturas y actividades por categoría y muestra completados, activos y vacíos."),
        ("Progreso: ", "se alimenta de intentos y porcentajes; se consume en feed, ruta, perfil y panel docente."),
        ("Logros: ", "buildLogros calcula logros y rangos a partir de progreso y publicaciones."),
        ("Panel docente: ", "overview, promedio, intentos, abandono, rendimiento por área, timeline, riesgo y tabla de estudiantes."),
    ])
    add_heading(doc, "5.5 Perfil", 2)
    add_bullets(doc, [
        ("Identidad: ", "nombre, rol, grado y avatar representado actualmente por inicial."),
        ("Avance: ", "porcentaje general y barras por tema."),
        ("Publicaciones propias: ", "aprobadas, en revisión y que requieren cambios."),
        ("Cuenta: ", "enlace de cierre de sesión; la API contempla actualizar perfil y avatar."),
    ])

    add_heading(doc, "6. Motor unificado de quizzes y juego en vivo", 1)
    add_heading(doc, "6.1 Constructor docente", 2)
    add_para(doc, "TeacherCreatorPage concentra la creación en una operación atómica hacia /api/quizzes. El docente define metadatos, reglas y una secuencia de bloques. Los medios se pueden cargar antes y quedan referenciados en el payload final.")
    add_table(doc, ["Bloque", "Uso", "Califica", "Media"], [
        ("single_choice", "Una respuesta correcta entre opciones.", "Sí", "Imagen/audio en bloque u opción"),
        ("multiple_choice", "Varias respuestas correctas.", "Sí", "Imagen/audio en bloque u opción"),
        ("true_false", "Afirmación verdadera o falsa.", "Sí", "Media general"),
        ("short_text", "Respuesta breve escrita.", "Sí, según respuesta configurada", "Media general"),
        ("flashcard", "Ficha de estudio o repaso.", "No necesariamente", "Imagen/audio"),
        ("info", "Bloque informativo dentro del recorrido.", "No", "Imagen/audio"),
    ], [1900, 3600, 1200, 2660])
    add_bullets(doc, [
        ("Reglas: ", "intentos máximos, tiempo global, porcentaje de aprobación, barajar bloques, barajar opciones, feedback y progreso."),
        ("Validación: ", "campos obligatorios, opciones, respuestas, puntaje, tiempo y coherencia de cada tipo."),
        ("Orden: ", "flechas para mover bloques y selector para cambiar el tipo sin perder el contenido común."),
        ("Persistencia: ", "quizzes, items, opciones, medios y versiones se guardan en tablas separadas."),
    ])
    add_heading(doc, "6.2 Reproductor unificado", 2)
    add_bullets(doc, [
        ("UnifiedQuiz: ", "renderiza todos los bloques con una sola experiencia y controla progreso, respuestas y resultado."),
        ("GameEngine: ", "actúa como punto de entrada y delega al motor unificado; ya no existe un quiz estándar separado como experiencia principal."),
        ("Actividad legacy: ", "ActivityDetailPage intenta cargar el quiz nuevo y conserva una ruta de compatibilidad para actividades antiguas con preguntas existentes."),
        ("Intentos: ", "se consulta el historial del estudiante antes de iniciar y se informa cuando se agotaron los intentos."),
        ("Resultado: ", "se registra puntaje, respuestas, tiempo, estado y resultado de aprobación."),
    ])
    add_heading(doc, "6.3 Juego multijugador", 2)
    add_bullets(doc, [
        ("Sala: ", "código alfanumérico, actividad asociada, máximo de jugadores y estados esperando/activa/finalizada."),
        ("QR: ", "backend genera QR o devuelve URL de unión; GameRoomPage lo muestra al docente."),
        ("WebSocket: ", "join, start, answer, next; emite room_state, player_joined, game_start, answer_ack, question_results y game_over."),
        ("Estudiante: ", "introduce código, espera en lobby, responde desde el móvil y recibe confirmación."),
        ("Docente: ", "ve jugadores, inicia partida, avanza preguntas y proyecta resultados."),
    ])
    add_callout(doc, "Brecha funcional del juego en vivo", "El backend sí tiene POST /api/rooms, listado, inicio, cierre y resultados, pero no hay una pantalla de administración que permita seleccionar una actividad, crear la sala y navegar automáticamente a /admin/sala/:code. Esa es una de las mejoras de mayor impacto para que la capacidad sea realmente descubrible.", fill=LIGHT_RED, accent="B91C1C")

    add_heading(doc, "7. Catálogo de API y contratos principales", 1)
    add_para(doc, "El backend registra rutas por dominio en api/routes/__init__.py. El siguiente catálogo resume los contratos disponibles, no sustituye la documentación OpenAPI de FastAPI.")
    add_table(doc, ["Dominio", "Endpoints principales", "Consumidor visible", "Estado funcional"], [
        ("Auth", "POST /auth/login, POST /auth/logout, GET /auth/me", "Login, AuthContext", "Completo"),
        ("Usuarios", "CRUD /users, /users/me, avatar, enroll, cursos", "Perfil solo parcialmente; sin gestión administrativa", "Parcial"),
        ("Cursos", "CRUD /courses, estudiantes por curso", "Ninguna pantalla dedicada", "Pendiente de UI"),
        ("Categorías", "CRUD /categories", "DataContext y ExplorePage", "Completo para consumo; administración pendiente"),
        ("Publicaciones", "CRUD, slug, portada, moderate, history", "Feed, crear, moderar, lecturas", "Completo"),
        ("Preguntas legacy", "CRUD /questions y listado por actividad", "Compatibilidad de actividades", "Mantenimiento/legacy"),
        ("Quizzes", "CRUD, publish, media, versions, quiz por actividad", "Creador, launcher, UnifiedQuiz", "Completo"),
        ("Intentos", "crear, listar, consultar, enviar, progreso", "ActivityDetail, resultados, dashboard", "Completo"),
        ("Recursos", "CRUD, aprobar, vínculos a publicaciones/actividades", "ResourcesPage; gestión no visible", "Parcial"),
        ("Comentarios", "crear, pendientes, moderar", "Detalle muestra formulario simulado", "Parcial"),
        ("Eventos", "crear y listar eventos de aprendizaje", "No se ve uso directo en pantallas", "Pendiente de conexión"),
        ("Auditoría", "crear y listar entradas", "Sin pantalla visible; seguridad fuera de alcance", "Pendiente de UI"),
        ("Stats", "overview, students, timeline, activities", "AdminDashboard", "Completo"),
        ("Rooms", "crear, listar, estado, QR, start, resultados, delete", "GameRoom/GameJoin sin creador de sala", "Parcial"),
    ], [1500, 3600, 2500, 1760], font_size=8.5)

    add_heading(doc, "8. Modelo de datos y relaciones", 1)
    add_para(doc, "La persistencia es SQLite con WAL, foreign keys y busy timeout configurados en schema.sql. Las entidades se dividen entre organización escolar, contenido social, aprendizaje evaluable, recursos, telemetría y juego en vivo.")
    add_table(doc, ["Grupo", "Tablas", "Relación funcional"], [
        ("Organización", "cursos, usuarios, inscripciones, sesiones", "Usuarios pertenecen a cursos y pueden tener rol de estudiante o docente por curso."),
        ("Taxonomía", "categorías, etiquetas", "Categorías agrupan contenidos por área; etiquetas clasifican publicaciones."),
        ("Contenido social", "publicaciones, publicacion_etiquetas, moderaciones, comentarios", "El contenido pasa por estados y puede recibir comentarios moderados."),
        ("Contenidos educativos", "actividades, preguntas, recursos, vínculos de recursos", "Lecturas, actividades y recursos se organizan por área/categoría."),
        ("Evaluación", "intentos, respuestas_intento, quizzes, quiz_items, quiz_opciones", "El modelo nuevo convive con preguntas legacy y registra resultados por intento."),
        ("Multimedia/versionado", "quiz_media_assets, quiz_versions", "Medios y snapshots permiten enriquecer y versionar quizzes."),
        ("Telemetría", "eventos_aprendizaje, auditoria", "Registra actividad y acciones administrativas; falta hacerla visible en producto."),
        ("Juego en vivo", "rooms, game_sessions, achievements, user_achievements", "Salas y sesiones alimentan podio y posibles logros."),
    ], [1800, 3000, 4560])
    add_heading(doc, "8.1 Tablas completas del esquema", 2)
    add_table(doc, ["Tabla", "Propósito", "Claves o notas relevantes"], [
        ("cursos", "Cursos escolares", "grado, jornada, año, activo"),
        ("usuarios", "Identidad y rol", "codigo único, correo, rol, activo"),
        ("inscripciones", "Membresía usuario-curso", "rol_en_curso y vigencia"),
        ("sesiones", "Sesiones persistidas", "token_hash, expira_en, activa"),
        ("categorias", "Áreas temáticas", "slug y area controlada"),
        ("etiquetas", "Etiquetas reutilizables", "nombre y slug únicos"),
        ("publicaciones", "Posts y lecturas", "estado, visibilidad, portada, moderador"),
        ("publicacion_etiquetas", "N:M publicaciones-etiquetas", "PK compuesta"),
        ("moderaciones", "Historial de decisiones", "publicación, docente, decisión"),
        ("comentarios", "Comentarios y respuestas", "árbol por comentario_padre_id"),
        ("actividades", "Actividades legacy y base de quiz", "área, dificultad, intentos máximos"),
        ("preguntas", "Preguntas legacy", "tipo, opciones, correcta, orden"),
        ("intentos", "Intentos de estudiante", "estado, porcentaje, tiempo"),
        ("respuestas_intento", "Respuestas legacy", "respuesta, correcta, puntaje"),
        ("recursos", "PDF, imagen, video y URL", "estado de aprobación y área"),
        ("publicacion_recursos / actividad_recursos", "Vínculos de recursos", "relaciones N:M"),
        ("eventos_aprendizaje", "Eventos de uso/aprendizaje", "actor, sesión y objeto"),
        ("auditoria", "Registro de acciones", "usuario, entidad, acción"),
        ("rooms / game_sessions", "Salas y partidas", "código, actividad, resultados"),
        ("achievements / user_achievements", "Logros", "catálogo y obtención por usuario"),
        ("quizzes", "Metadatos del motor nuevo", "estado, creador, configuración"),
        ("quiz_items", "Bloques del quiz", "tipo, orden, puntaje, configuración"),
        ("quiz_opciones", "Opciones de bloques", "valor, correcta, feedback"),
        ("quiz_media_assets", "Imágenes, audios y archivos", "asociable a quiz/item/opción"),
        ("quiz_versions", "Snapshots", "número de versión y JSON"),
        ("quiz_respuestas_intento", "Respuestas del motor nuevo", "respuesta, puntaje, tiempo"),
        ("quiz_attempt_events", "Eventos de reproducción", "navegación y eventos por item"),
    ], [2300, 3000, 4060], font_size=8.6)

    add_heading(doc, "9. Estado de completitud y brechas", 1)
    add_para(doc, "Esta matriz es la parte más importante para decidir qué agregar. “Parcial” significa que hay una experiencia o un contrato, pero no el circuito completo para el usuario objetivo.")
    add_table(doc, ["Funcionalidad", "Estado", "Evidencia actual", "Qué falta para cerrarla"], [
        ("Inicio estudiante", "Completo", "Feed, retos, filtros, publicaciones y creación.", "Pruebas de contenido y pulido visual; no es una brecha estructural."),
        ("Inicio docente", "Completo", "Dashboard separado con KPIs y acciones docentes.", "Integrar accesos directos a salas, quizzes propios y tareas del día."),
        ("Crear publicación con imagen", "Completo", "Preview, validación, subida y refresco de datos.", "Mostrar estado de imagen de forma consistente en todas las vistas."),
        ("Ver imagen publicada", "Completo condicionado", "Detalle usa imagen_portada y feed transforma imagen.", "Verificar datos antiguos y fallbacks cuando la URL no esté disponible."),
        ("Comentarios", "Parcial", "Backend listo; el detalle muestra alert() en vez de api.createComment.", "Persistir, listar, actualizar contexto y moderar desde la UI."),
        ("Quiz unificado", "Completo", "Creador, tablas nuevas y UnifiedQuiz.", "Añadir edición/listado/duplicado y analítica docente avanzada."),
        ("Quiz estándar antiguo", "Retirado", "GameEngine delega al motor unificado; legacy solo sirve compatibilidad.", "Eliminar progresivamente datos/pantallas legacy cuando no haya dependencias."),
        ("Catálogo docente de quizzes", "Parcial", "Launcher lista actividades disponibles.", "Vista de mis quizzes, borradores, publicar, archivar, editar y versionar."),
        ("Salas en vivo", "Parcial", "API, QR, WebSocket y vistas de sala/unión.", "UI para crear sala, elegir actividad, listar salas y cerrar sesión."),
        ("Cursos y matrículas", "Pendiente de UI", "Schema y endpoints existen.", "Pantallas para crear curso, asignar docentes/estudiantes y cambiar estado."),
        ("Gestión de usuarios", "Pendiente de UI", "CRUD y perfil backend.", "Backoffice con búsqueda, edición, roles y estado activo."),
        ("Recursos docentes", "Parcial", "Consulta y vínculos disponibles.", "Subir, aprobar, editar y eliminar desde una biblioteca docente."),
        ("Eventos de aprendizaje", "Pendiente de conexión", "API y tabla existen.", "Instrumentar eventos relevantes y mostrar métricas útiles."),
        ("Auditoría funcional", "Pendiente de UI", "API y tabla existen.", "Pantalla administrativa de historial, filtros y exportación."),
        ("Notificaciones", "Pendiente", "TopBar muestra accesos, pero no hay centro de notificaciones persistente.", "Diseñar notificaciones para cambios, feedback, intentos y pendientes."),
        ("Búsqueda global", "Parcial", "Búsqueda existe dentro de Explore y Lecturas.", "Unificar búsqueda de publicaciones, actividades y recursos con estados vacíos."),
    ], [1900, 1200, 3000, 3260], font_size=8.3)

    add_heading(doc, "10. Backlog funcional recomendado", 1)
    add_heading(doc, "Prioridad P0: cerrar circuitos ya iniciados", 2)
    add_numbered(doc, [
        "Crear la pantalla docente de salas: seleccionar actividad, indicar máximo de jugadores, crear código, mostrar QR y entrar automáticamente a la sala.",
        "Conectar comentarios reales: POST, lista actualizada, estado pendiente y moderación desde el panel docente.",
        "Crear la biblioteca de quizzes del docente: mis quizzes, borradores, publicados, editar, duplicar, archivar y ver versiones.",
        "Añadir biblioteca de recursos para docentes: cargar, vincular a publicación/actividad, aprobar y eliminar.",
        "Revisar las rutas compartidas en contexto docente para que títulos, acciones y copy no parezcan de estudiante.",
    ])
    add_heading(doc, "Prioridad P1: ampliar capacidad de enseñanza", 2)
    add_numbered(doc, [
        "Administración de cursos y matrículas con selección de curso activo para cada docente.",
        "Gestión de usuarios y perfiles docentes diferenciados del perfil de estudiante.",
        "Asignación de quizzes/actividades a cursos o grupos, con fechas, disponibilidad y límites.",
        "Centro de resultados por quiz: distribución de puntajes, preguntas difíciles, intentos y estudiantes sin actividad.",
        "Notificaciones de moderación, feedback, nuevas publicaciones y actividades asignadas.",
    ])
    add_heading(doc, "Prioridad P2: completar producto educativo", 2)
    add_numbered(doc, [
        "Ruta de aprendizaje configurable por curso, con estados reales de bloqueado, disponible, actual y completado.",
        "Contenido multimedia más rico en lecturas y recursos, con vista previa y metadatos educativos.",
        "Plantillas de quiz reutilizables y duplicación por parte del docente.",
        "Logros y gamificación conectados a eventos reales, no solo al cálculo de progreso.",
        "Exportes funcionales de progreso y resultados para el docente.",
    ])
    add_callout(doc, "Lo que no se debe añadir todavía", "No conviene crear más tipos de quiz, más pantallas de navegación o más entidades de contenido antes de cerrar las brechas de descubribilidad: crear salas, administrar quizzes propios, persistir comentarios y gestionar recursos. La base técnica ya permite crecer; primero hay que completar los circuitos visibles.", fill=LIGHT_YELLOW, accent="A16207")

    add_heading(doc, "11. Criterios de aceptación de una aplicación funcionalmente completa", 1)
    add_para(doc, "Sin considerar seguridad, el producto puede considerarse funcionalmente completo cuando una persona de cada rol pueda realizar el siguiente recorrido sin rutas manuales, errores silenciosos ni intervención directa en la base de datos:")
    add_heading(doc, "Estudiante", 2)
    add_bullets(doc, [
        "Iniciar sesión y llegar a su inicio correcto.",
        "Explorar un tema, abrir una lectura con imagen y consultar recursos vinculados.",
        "Realizar un quiz de cada tipo soportado y ver el resultado.",
        "Ver su progreso y que la ruta cambie después de completar una actividad.",
        "Crear una publicación con imagen, recibir feedback y ver el estado de revisión.",
        "Comentar una lectura y ver el comentario después de ser aprobado.",
        "Unirse a una sala usando QR o código y participar en una partida.",
    ])
    add_heading(doc, "Docente", 2)
    add_bullets(doc, [
        "Entrar a un panel docente separado del feed del estudiante.",
        "Crear, editar, previsualizar, publicar, duplicar y archivar un quiz.",
        "Subir imágenes/audio y comprobar que se muestran en el reproductor.",
        "Crear una sala seleccionando actividad, mostrar QR, iniciar, avanzar y cerrar.",
        "Revisar una publicación, pedir cambios y comprobar el nuevo estado del estudiante.",
        "Consultar resultados por estudiante, actividad y área con acciones de acompañamiento.",
        "Subir y vincular recursos a lecturas o actividades.",
    ])
    add_heading(doc, "Administrador", 2)
    add_bullets(doc, [
        "Crear y administrar cursos, usuarios, matrículas y categorías desde una interfaz dedicada.",
        "Consultar actividad del sistema, auditoría funcional y datos agregados.",
        "Resolver incidencias de contenido sin editar directamente la base de datos.",
    ])

    add_heading(doc, "12. Conclusión", 1)
    add_para(doc, "La aplicación no está vacía ni necesita una reescritura general. Tiene una arquitectura coherente para el aprendizaje, un motor de quizzes nuevo y una separación de roles que ya permite trabajar con estudiantes y docentes. La decisión de producto más importante ahora es pasar de “capacidades disponibles” a “flujos completos y descubribles”.")
    add_para(doc, "La siguiente iteración debería concentrarse en el panel docente: biblioteca de quizzes, creación de salas, recursos, comentarios y gestión de cursos. Después de esos cierres, será más fácil evaluar qué características educativas nuevas aportan valor real.")
    add_callout(doc, "Siguiente paso recomendado", "Usar la matriz de la sección 9 como tablero de trabajo y cerrar primero las funcionalidades marcadas como Parcial que ya tienen backend: salas, comentarios, biblioteca de quizzes y recursos docentes. Luego validar los recorridos de aceptación con un usuario estudiante y uno docente.", fill=LIGHT_GREEN, accent="15803D")

    # Add a compact source note without turning this into a security report.
    add_para(doc, "Fuentes revisadas: frontend/src/App.jsx, layouts y páginas React; frontend/src/services/api.js; api/app.py; api/routes; api/services; schema.sql; start-debug.bat; logs y documentación local. Fecha de corte: 1 de agosto de 2026.", italic=True, color=MUTED, size=9, after=0)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
